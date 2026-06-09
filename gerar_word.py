#!/usr/bin/env python3
"""
Gera dois arquivos Word (.docx) dos cartões de voltas:
  cartoes-voltas.docx  — branco com versículos
  cartoes-azul.docx    — azul escuro com números grandes
"""
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

laps = [
    (1,  "3,35 km",  "O meu socorro vem do Senhor, que fez o céu e a terra.",
     "My help comes from the Lord, the Maker of heaven and earth.", "Sl / Ps 121:2"),
    (2,  "6,70 km",  "O Senhor é a minha força e o meu escudo; nele o meu coração confia.",
     "The Lord is my strength and my shield; my heart trusts in him.", "Sl / Ps 28:7"),
    (3,  "10,05 km", "Tudo posso naquele que me fortalece.",
     "I can do all things through Christ who strengthens me.", "Fp / Ph 4:13"),
    (4,  "13,40 km", "Entrega o teu caminho ao Senhor; confia nele, e ele tudo fará.",
     "Commit your way to the Lord; trust in him and he will do this.", "Sl / Ps 37:5"),
    (5,  "16,75 km", "Seja forte e corajoso! Não se apavore, pois o Senhor, o teu Deus, estará contigo.",
     "Be strong and courageous. Do not be afraid; the Lord your God will be with you.", "Js / Jo 1:9"),
    (6,  "20,10 km", "Os que esperam no Senhor renovam as suas forças. Voam alto como águias.",
     "Those who hope in the Lord will renew their strength. They will soar on wings like eagles.", "Is / Is 40:31"),
    (7,  "23,45 km", "Deus é o nosso refúgio e a nossa força, auxílio sempre presente na hora da angústia.",
     "God is our refuge and strength, an ever-present help in trouble.", "Sl / Ps 46:1"),
    (8,  "26,80 km", "Em todas estas coisas somos mais que vencedores por meio daquele que nos amou.",
     "In all these things we are more than conquerors through him who loved us.", "Rm / Ro 8:37"),
    (9,  "30,15 km", "A minha graça é suficiente para você, pois o meu poder se aperfeiçoa na fraqueza.",
     "My grace is sufficient for you, for my power is made perfect in weakness.", "2Co / 2Co 12:9"),
    (10, "33,50 km", "Ainda que eu ande pelo vale da sombra da morte, não temerei mal nenhum, pois tu estás comigo.",
     "Even though I walk through the darkest valley, I will fear no evil, for you are with me.", "Sl / Ps 23:4"),
    (11, "36,85 km", "Sede fortes e corajosos. Não temais, pois o Senhor, teu Deus, vai contigo.",
     "Be strong and courageous. Do not be afraid; the Lord your God goes with you.", "Dt / Dt 31:6"),
    (12, "40,20 km", "Sede fortalecidos no Senhor e na força do seu poder.",
     "Be strong in the Lord and in his mighty power.", "Ef / Ep 6:10"),
    (13, "43,55 km", "O Senhor é a minha rocha, a minha fortaleza e o meu libertador.",
     "The Lord is my rock, my fortress and my deliverer.", "Sl / Ps 18:2"),
    (14, "46,90 km", "Não temas, porque eu sou contigo; não te assombres, porque eu sou o teu Deus.",
     "Do not fear, for I am with you; do not be dismayed, for I am your God.", "Is / Is 41:10"),
    (15, "50,25 km", "Corramos com perseverança a corrida que nos está proposta.",
     "Let us run with perseverance the race marked out for us.", "Hb / He 12:1"),
    (16, "53,60 km", "O Senhor dará força ao seu povo; o Senhor abençoará o seu povo com a paz.",
     "The Lord gives strength to his people; the Lord blesses his people with peace.", "Sl / Ps 29:11"),
    (17, "56,95 km", "Graças a Deus que nos dá a vitória por meio de nosso Senhor Jesus Cristo.",
     "Thanks be to God! He gives us the victory through our Lord Jesus Christ.", "1Co / 1Co 15:57"),
    (18, "60,30 km", "No mundo tereis aflições; mas tende bom ânimo, eu venci o mundo.",
     "In this world you will have trouble. But take heart! I have overcome the world.", "Jo / Jn 16:33"),
    (19, "64 km",    "O vencedor herdará estas coisas, e eu serei o seu Deus e ele será o meu filho.",
     "The one who is victorious will inherit all this, and I will be his God and he will be my son.", "Ap / Rv 21:7"),
]

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='000000', sz='12'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), sz)
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_para(cell, text, bold=False, italic=False, size=12,
             color=None, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# 1 — WHITE + VERSES
# ══════════════════════════════════════════════════════════════════════════════
def make_cartolina_docx(path):
    doc = Document()

    # Page setup: A4, no margins
    sec = doc.sections[0]
    sec.page_width  = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin    = Mm(0)
    sec.bottom_margin = Mm(0)
    sec.left_margin   = Mm(0)
    sec.right_margin  = Mm(0)

    for idx, (num, km, pt, en, ref) in enumerate(laps):
        if idx > 0 and idx % 3 == 0:
            doc.add_page_break()

        # Card via 1-column table
        card_h = Mm(99) if num < 19 else Mm(297)
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'

        cell = tbl.rows[0].cells[0]
        cell.width = Mm(210)

        # set row height
        tr = tbl.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(card_h.pt * 20)))
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, 'FFFFFF')
        set_cell_border(cell, color='111111', sz='8')

        # clear default paragraph
        cell.paragraphs[0].clear()

        # VOLTA label
        add_para(cell, f'VOLTA', bold=True, size=7,
                 color=(180,180,180), align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)

        # Number
        num_size = 72 if num < 19 else 100
        num_color = (28, 132, 73) if num == 19 else (17, 17, 17)
        add_para(cell, str(num), bold=True, size=num_size, color=num_color, space_after=2)

        # km
        km_size = 18 if num == 19 else 13
        add_para(cell, km + (' — FINISH! 🏆' if num == 19 else ''),
                 bold=True, size=km_size, color=(28,132,73) if num==19 else (100,100,100),
                 space_after=6)

        # divider (em-dashes as separator)
        add_para(cell, '─' * 30, size=8, color=(180,180,180), space_after=4)

        # PT verse
        add_para(cell, pt, bold=True, size=11 if num==19 else 10,
                 color=(17,17,17), space_after=3)

        # EN verse
        add_para(cell, en, italic=True, size=9,
                 color=(100,100,100), space_after=4)

        # Reference (right-aligned)
        add_para(cell, ref, bold=True, size=8, color=(160,160,160),
                 align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.save(path)
    print(f'✅  Saved: {path}')

# ══════════════════════════════════════════════════════════════════════════════
# 2 — DARK BLUE
# ══════════════════════════════════════════════════════════════════════════════
def make_azul_docx(path):
    doc = Document()

    sec = doc.sections[0]
    sec.page_width  = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin    = Mm(0)
    sec.bottom_margin = Mm(0)
    sec.left_margin   = Mm(0)
    sec.right_margin  = Mm(0)

    for idx, (num, km, *_) in enumerate(laps):
        if idx > 0 and idx % 3 == 0:
            doc.add_page_break()

        card_h = Mm(99) if num < 19 else Mm(297)
        bg_color = '071050' if num == 19 else '0d1b6e'

        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = 'Table Grid'

        cell = tbl.rows[0].cells[0]
        cell.width = Mm(210)

        tr = tbl.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(card_h.pt * 20)))
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, bg_color)
        set_cell_border(cell, color='FFFFFF', sz='12')

        cell.paragraphs[0].clear()

        num_size = 100 if num == 19 else 80
        add_para(cell, str(num), bold=True, size=num_size,
                 color=(255,255,255), space_after=6)

        km_text = km + ' 🏆' if num == 19 else km
        add_para(cell, km_text, bold=True, size=18,
                 color=(255,255,255))

    doc.save(path)
    print(f'✅  Saved: {path}')


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    make_cartolina_docx(os.path.join(base, 'cartoes-voltas.docx'))
    make_azul_docx(os.path.join(base, 'cartoes-azul.docx'))
    print('\nDone! Open the .docx files in the desafio-64km folder.')
